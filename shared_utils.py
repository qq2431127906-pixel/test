# shared_utils.py —— 共享工具模块
import io
import os
import re
import sys
import streamlit as st

# ============================================================
# 0. LLM 统一封装
# ============================================================
_ENV_LOADED = False

def _load_env_file_fallback(env_path: str):
    """在 python-dotenv 不可用时读取简单的 KEY=VALUE 配置。"""
    with open(env_path, "r", encoding="utf-8-sig") as f:
        for raw_line in f:
            line = raw_line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, value = line.split("=", 1)
            key = key.strip()
            value = value.strip().strip('"').strip("'")
            if key and key not in os.environ:
                os.environ[key] = value

def load_app_env():
    """统一加载 .env；不覆盖系统环境变量，兼容源码与 PyInstaller 运行。"""
    global _ENV_LOADED
    if _ENV_LOADED:
        return

    try:
        from dotenv import load_dotenv
    except ImportError:
        load_dotenv = None

    candidate_dirs = []
    if getattr(sys, "frozen", False):
        candidate_dirs.append(os.path.dirname(os.path.abspath(sys.executable)))
    candidate_dirs.append(os.path.dirname(os.path.abspath(__file__)))
    candidate_dirs.append(os.getcwd())

    seen = set()
    for directory in candidate_dirs:
        env_path = os.path.join(directory, ".env")
        if env_path in seen:
            continue
        seen.add(env_path)
        if os.path.exists(env_path):
            if load_dotenv:
                load_dotenv(env_path, override=False)
            else:
                _load_env_file_fallback(env_path)
            break

    _ENV_LOADED = True

def _get_dashscope_llm(api_key, max_tokens, temperature):
    """从 llama_index 或 dashscope 获取 LLM，兼容两种安装方式"""
    try:
        from llama_index.llms.dashscope import DashScope
        return DashScope(model_name="qwen-turbo", api_key=api_key,
                         max_tokens=max_tokens, temperature=temperature)
    except ImportError as llama_index_error:
        # 直接用 dashscope SDK，不依赖 llama_index
        try:
            from dashscope import Generation
        except ImportError as dashscope_error:
            raise RuntimeError(
                "缺少 DashScope 依赖。请运行：pip install -r requirements.txt "
                "或单独安装：pip install dashscope llama-index-llms-dashscope"
            ) from dashscope_error

        class _DashScopeCompatible:
            def __init__(self, model_name, api_key, max_tokens, temperature):
                self.model_name = model_name
                self.api_key = api_key
                self.max_tokens = max_tokens
                self.temperature = temperature
            def complete(self, prompt):
                resp = Generation.call(
                    model=self.model_name,
                    prompt=str(prompt),
                    api_key=self.api_key,
                    max_tokens=self.max_tokens,
                    temperature=self.temperature,
                )
                if resp.status_code != 200:
                    raise RuntimeError(f"DashScope error: {resp.message}")
                class _Response:
                    text = resp.output.get("text", "")
                return _Response()
        return _DashScopeCompatible(model_name="qwen-turbo", api_key=api_key,
                                     max_tokens=max_tokens, temperature=temperature)

def get_dashscope_api_key() -> str:
    load_app_env()
    key = os.getenv("DASHSCOPE_API_KEY")
    if not key or key.strip() in {"", "sk-你的API密钥"}:
        raise RuntimeError("请先在 .env 文件中配置 DASHSCOPE_API_KEY 后再使用 AI 生成功能。")
    return key.strip()

def get_llm(max_tokens=8192, temperature=0.1):
    key = get_dashscope_api_key()
    return _get_dashscope_llm(key, max_tokens, temperature)

def format_ai_error(error: Exception) -> str:
    """把常见模型/API异常转换成普通用户能看懂的提示。"""
    msg = str(error).strip() or error.__class__.__name__
    lower = msg.lower()
    if "dashscope_api_key" in lower or "api key" in lower or "401" in lower or "unauthorized" in lower:
        return "请检查 .env 文件中的 DASHSCOPE_API_KEY 是否已填写且有效。"
    if "403" in lower or "forbidden" in lower or "permission" in lower or "qwen-vl-max" in lower:
        return "当前 API Key 可能没有对应模型权限，请在 DashScope 控制台确认 qwen-turbo / qwen-vl-max 权限。"
    if "quota" in lower or "insufficient" in lower or "balance" in lower or "余额" in msg or "额度" in msg:
        return "DashScope 额度不足或账户欠费，请检查账号余额和调用额度。"
    if "timeout" in lower or "timed out" in lower:
        return "请求模型超时，请稍后重试，或检查网络连接。"
    if "connection" in lower or "network" in lower or "proxy" in lower:
        return "网络连接失败，请检查网络、代理或防火墙设置。"
    return f"AI 服务调用失败：{msg}"

def show_ai_error(error: Exception, prefix: str = "AI 生成失败"):
    st.error(f"{prefix}：{format_ai_error(error)}")

# ============================================================
# 1. 依赖检测
# ============================================================
def check_dependencies():
    """返回 (ok: bool, missing: list[str]) —— 缺少的包名列表"""
    missing = []
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        missing.append("python-docx")
    try:
        import pdfplumber  # noqa: F401
    except ImportError:
        missing.append("pdfplumber")
    try:
        from pypdf import PdfReader  # noqa: F401
    except ImportError:
        missing.append("pypdf")
    return len(missing) == 0, missing

def show_dependency_install_hint(missing):
    """在页面上显示安装提示"""
    if not missing:
        return
    pkgs = " ".join(missing)
    st.warning(
        f"缺少文档解析依赖：`{', '.join(missing)}`。"
        f"请在终端运行以下命令安装：\n\n"
        f"```bash\npip install {pkgs}\n```\n\n"
        f"安装后刷新页面即可使用 PDF / Word 上传功能。"
    )

# ============================================================
# 2. 文档正文抽取
# ============================================================
def extract_docx_text(uploaded_file) -> str:
    """从 .docx 文件抽取正文（段落 + 表格）"""
    from docx import Document
    doc = Document(io.BytesIO(uploaded_file.getvalue()))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)

def extract_pdf_text(uploaded_file, max_pages=30) -> str:
    """从 PDF 抽取正文；优先 pdfplumber，回退 pypdf"""
    data = uploaded_file.getvalue()
    # 优先 pdfplumber
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages[:max_pages]:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        result = "\n".join(text_parts).strip()
        if result:
            return result
    except Exception:
        pass
    # 回退 pypdf
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    text_parts = []
    for page in reader.pages[:max_pages]:
        t = page.extract_text()
        if t:
            text_parts.append(t)
    return "\n".join(text_parts).strip()

def extract_text_from_uploaded_file(uploaded_file) -> str:
    """
    统一入口：根据扩展名调用对应抽取逻辑。
    返回抽取到的纯文本；.doc 文件抛出 RuntimeError 提示转换。
    """
    name = uploaded_file.name.lower()
    if name.endswith(".docx"):
        return extract_docx_text(uploaded_file)
    if name.endswith(".pdf"):
        return extract_pdf_text(uploaded_file)
    if name.endswith(".doc"):
        raise RuntimeError("暂不支持旧版 .doc 文件，请在 Word/WPS 中另存为 .docx 后上传。")
    raise RuntimeError("仅支持 PDF 和 DOCX 文件。")

def extract_texts_from_files(uploaded_files) -> list[dict]:
    """
    批量抽取，返回列表，每项：
    {
        "name": 文件名,
        "text": 抽取正文,
        "error": None | str,
        "size": 字符数
    }
    """
    results = []
    for f in uploaded_files:
        item = {"name": f.name, "text": "", "error": None, "size": 0}
        try:
            text = extract_text_from_uploaded_file(f)
            if not text or not text.strip():
                item["error"] = "未能抽取到正文文本，可能是扫描件或图片型 PDF。"
            else:
                item["text"] = text
                item["size"] = len(text)
        except RuntimeError as e:
            item["error"] = str(e)
        except Exception as e:
            item["error"] = f"解析异常：{e}"
        results.append(item)
    return results

# ============================================================
# 3. 长文本分段
# ============================================================
def chunk_text_by_paragraphs(text: str, max_chars: int = 4000) -> list[str]:
    """
    按段落边界拆分文本，每段不超过 max_chars。
    尽量保持自然段完整。
    """
    paragraphs = re.split(r"\n\s*\n", text)
    chunks = []
    current = ""
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current) + len(para) + 2 <= max_chars:
            current = (current + "\n\n" + para).strip()
        else:
            if current:
                chunks.append(current)
            # 如果单个段落就超过限制，强制按字符切
            if len(para) > max_chars:
                for i in range(0, len(para), max_chars):
                    chunks.append(para[i:i + max_chars])
                current = ""
            else:
                current = para
    if current:
        chunks.append(current)
    return chunks

# ============================================================
# 4. Markdown 下载按钮
# ============================================================
def render_download_button(content: str, filename: str, label: str = "📥 下载完整报告"):
    """在页面上放一个 Markdown 下载按钮"""
    if not content:
        return
    st.download_button(
        label=label,
        data=content,
        file_name=filename,
        mime="text/markdown",
        use_container_width=True,
    )
