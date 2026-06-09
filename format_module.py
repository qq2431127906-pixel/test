import io
import os
import re

os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"

import streamlit as st
from shared_utils import get_llm
from ui_theme import render_page_shell
from dotenv import load_dotenv
load_dotenv()

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None


api_key = os.getenv("DASHSCOPE_API_KEY", "") or "sk-9e84da6799aa4022947b585b78e0fb31"


REFERENCE_FROM_DOCUMENT_PROMPT = """
你是专业的学术参考文献格式化助手。请根据用户上传文档中可识别的信息，
自动生成符合 GB/T 7714-2015 顺序编码制 的参考文献条目。

要求：
1. 优先识别作者、题名、文献类型、期刊/学校/会议/出版社、年份、卷期、页码、DOI、URL等信息。
2. 信息不足时，不要编造；缺失字段直接省略或采用该文献类型允许的最简格式。
3. 作者超过3人时，只保留前3人，后面加“等”。
4. 标点符号使用英文半角；中文姓名之间用英文逗号。
5. 每个上传文件生成1条参考文献，按 [1]、[2] 顺序输出。
6. 只输出参考文献条目，不要解释。

上传文档信息：
{document_context}
"""


CHINESE_NUMBERS = {
    "零": 0, "〇": 0, "一": 1, "二": 2, "两": 2, "三": 3, "四": 4,
    "五": 5, "六": 6, "七": 7, "八": 8, "九": 9
}


def chinese_to_int(value):
    value = value.strip()
    if value.isdigit():
        return int(value)
    if value in CHINESE_NUMBERS:
        return CHINESE_NUMBERS[value]
    if "十" in value:
        left, _, right = value.partition("十")
        tens = CHINESE_NUMBERS.get(left, 1) if left else 1
        ones = CHINESE_NUMBERS.get(right, 0) if right else 0
        return tens * 10 + ones
    return None


def get_heading_style_level(style_name):
    if not style_name:
        return None
    match = re.search(r"(?:Heading|标题)\s*([1-9])", style_name, re.I)
    if match:
        return int(match.group(1))
    return None


def extract_docx_blocks(uploaded_file):
    if Document is None:
        raise RuntimeError("缺少 python-docx 依赖，请先安装：pip install python-docx")

    doc = Document(io.BytesIO(uploaded_file.getvalue()))
    blocks = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append({
                "line": len(blocks) + 1,
                "text": text,
                "style": para.style.name if para.style else "",
                "source": "正文段落"
            })

    for table_index, table in enumerate(doc.tables, start=1):
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append({
                    "line": len(blocks) + 1,
                    "text": " | ".join(cells),
                    "style": "Table",
                    "source": f"表格{table_index}"
                })

    return blocks


def extract_pdf_text(uploaded_file, max_pages=8):
    data = uploaded_file.getvalue()

    if pdfplumber is not None:
        text_parts = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages[:max_pages]:
                text_parts.append(page.extract_text() or "")
        return "\n".join(text_parts).strip()

    if PdfReader is not None:
        reader = PdfReader(io.BytesIO(data))
        text_parts = []
        for page in reader.pages[:max_pages]:
            text_parts.append(page.extract_text() or "")
        return "\n".join(text_parts).strip()

    raise RuntimeError("缺少 PDF 解析依赖，请先安装：pip install pdfplumber pypdf")


def extract_text_from_uploaded_file(uploaded_file):
    name = uploaded_file.name.lower()

    if name.endswith(".docx"):
        blocks = extract_docx_blocks(uploaded_file)
        return "\n".join(block["text"] for block in blocks)

    if name.endswith(".pdf"):
        return extract_pdf_text(uploaded_file)

    if name.endswith(".doc"):
        raise RuntimeError("暂不支持旧版 .doc 文件，请在 Word/WPS 中另存为 .docx 后上传。")

    raise RuntimeError("仅支持 PDF 和 DOCX 文件。")


def build_reference_context(uploaded_files):
    chunks = []
    for index, file in enumerate(uploaded_files, start=1):
        try:
            text = extract_text_from_uploaded_file(file)
        except Exception as exc:
            chunks.append(f"文件{index}：{file.name}\n解析失败：{exc}")
            continue

        excerpt = text[:12000] if text else "未提取到正文文本"
        chunks.append(f"文件{index}：{file.name}\n提取内容：\n{excerpt}")

    return "\n\n---\n\n".join(chunks)


def generate_references_from_documents(uploaded_files):
    document_context = build_reference_context(uploaded_files)
    prompt = REFERENCE_FROM_DOCUMENT_PROMPT.format(document_context=document_context)
    return get_llm().complete(prompt).text.strip()


def parse_chapter_heading(text):
    match = re.match(r"^第([一二三四五六七八九十〇零两\d]+)章\s*(.+)$", text)
    if not match:
        return None
    return chinese_to_int(match.group(1))


def parse_numbered_heading(text):
    match = re.match(r"^(\d+(?:\.\d+){0,3})[\.、]?\s+(.+)$", text)
    if not match:
        return None
    numbers = [int(item) for item in match.group(1).split(".")]
    return numbers


def add_issue(issues, block, issue_type, message, suggestion):
    issues.append({
        "type": issue_type,
        "line": block["line"],
        "source": block.get("source", "正文段落"),
        "style": block.get("style", ""),
        "old": block["text"],
        "msg": message,
        "suggestion": suggestion,
    })


def check_heading_hierarchy(blocks):
    issues = []
    current_chapter = None
    last_by_prefix = {}
    last_style_level = 0

    for block in blocks:
        text = block["text"]
        style_level = get_heading_style_level(block.get("style", ""))
        chapter_num = parse_chapter_heading(text)
        numbered = parse_numbered_heading(text)

        if chapter_num is not None:
            if current_chapter is not None and chapter_num != current_chapter + 1:
                add_issue(
                    issues,
                    block,
                    "标题层级",
                    f"章节编号不连续：上一章为第{current_chapter}章，当前为第{chapter_num}章。",
                    f"建议检查是否遗漏章节，或改为“第{current_chapter + 1}章 ...”。"
                )
            current_chapter = chapter_num
            last_by_prefix = {(chapter_num,): 0}
            last_style_level = 1
            continue

        if numbered:
            level = len(numbered)
            prefix = tuple(numbered[:-1])
            current = numbered[-1]

            if level == 1:
                if current_chapter is not None and numbered[0] != current_chapter:
                    add_issue(
                        issues,
                        block,
                        "标题层级",
                        f"一级标题编号与当前章节不一致：当前章节为第{current_chapter}章，标题编号为{numbered[0]}。",
                        f"建议改为“{current_chapter} {text.split(maxsplit=1)[-1]}”或检查章节顺序。"
                    )
                current_chapter = numbered[0]
                last_by_prefix[(current_chapter,)] = 0
            else:
                parent = tuple(numbered[:-1])
                if current_chapter is not None and numbered[0] != current_chapter:
                    add_issue(
                        issues,
                        block,
                        "标题层级",
                        f"节标题所属章节错误：当前章节为第{current_chapter}章，但该标题属于第{numbered[0]}章。",
                        f"建议将编号调整为以 {current_chapter}. 开头，或移动到对应章节。"
                    )

                parent_seen = parent in last_by_prefix or level == 2
                if not parent_seen:
                    add_issue(
                        issues,
                        block,
                        "标题层级",
                        f"标题层级跳跃：出现 {'.'.join(map(str, numbered))} 前，未检测到上级标题 {'.'.join(map(str, parent))}。",
                        "建议补充上级标题，或降低该标题层级。"
                    )

            expected = last_by_prefix.get(prefix, 0) + 1
            if current != expected:
                add_issue(
                    issues,
                    block,
                    "标题编号",
                    f"同级标题编号不连续：当前为 {'.'.join(map(str, numbered))}，按前文顺序应为 {'.'.join(map(str, list(prefix) + [expected]))}。",
                    f"建议检查是否漏写标题，或改为 {'.'.join(map(str, list(prefix) + [expected]))}。"
                )

            last_by_prefix[prefix] = current
            last_by_prefix[tuple(numbered)] = 0
            last_style_level = max(last_style_level, level)

        if style_level is not None:
            if last_style_level and style_level > last_style_level + 1:
                add_issue(
                    issues,
                    block,
                    "Word样式层级",
                    f"Word标题样式从 {last_style_level} 级跳到 {style_level} 级。",
                    "建议按“标题1 -> 标题2 -> 标题3”的顺序设置样式，避免跳级。"
                )
            last_style_level = style_level

    return issues


def check_caption_sequence(blocks, label, issue_type):
    issues = []
    last_by_chapter = {}
    current_chapter = None
    correct_pattern = re.compile(rf"^{label}\s*(\d+)[\-\.](\d+)\s*\S+")

    for block in blocks:
        text = block["text"].strip()
        chapter_num = parse_chapter_heading(text)
        numbered = parse_numbered_heading(text)
        if chapter_num is not None:
            current_chapter = chapter_num
        elif numbered and len(numbered) == 1:
            current_chapter = numbered[0]

        if not text.startswith(label):
            continue

        match = correct_pattern.match(text)
        if not match:
            add_issue(
                issues,
                block,
                issue_type,
                f"{label}编号格式不规范。",
                f"建议使用“{label}X-Y 标题”格式，例如“{label}1-1 系统结构图”。"
            )
            continue

        chapter = int(match.group(1))
        number = int(match.group(2))

        if current_chapter is not None and chapter != current_chapter:
            add_issue(
                issues,
                block,
                issue_type,
                f"{label}编号章节号与当前位置不一致：当前章节为第{current_chapter}章，编号为{label}{chapter}-{number}。",
                f"建议改为“{label}{current_chapter}-{number} ...”或检查该{label}所在章节。"
            )

        expected = last_by_chapter.get(chapter, 0) + 1
        if number != expected:
            add_issue(
                issues,
                block,
                issue_type,
                f"{label}序号不连续：本章上一处为{label}{chapter}-{last_by_chapter.get(chapter, 0)}，当前为{label}{chapter}-{number}。",
                f"建议检查是否遗漏{label}，或改为“{label}{chapter}-{expected} ...”。"
            )

        last_by_chapter[chapter] = number

    return issues


def check_formula_numbering(blocks):
    issues = []
    last_by_chapter = {}
    current_chapter = None
    correct_pattern = re.compile(r"(?:式\s*)?[\(（](\d+)[\-\.](\d+)[\)）]")

    for block in blocks:
        text = block["text"].strip()
        chapter_num = parse_chapter_heading(text)
        numbered = parse_numbered_heading(text)
        if chapter_num is not None:
            current_chapter = chapter_num
        elif numbered and len(numbered) == 1:
            current_chapter = numbered[0]

        if "公式" not in text and "式" not in text and not re.search(r"[\(（]\d+[\)）]$", text):
            continue

        match = correct_pattern.search(text)
        if not match:
            add_issue(
                issues,
                block,
                "公式编号",
                "检测到公式相关内容，但未发现规范的公式编号。",
                "建议使用“式(章号-序号)”或“(章号-序号)”格式，例如“式(2-1)”。"
            )
            continue

        chapter = int(match.group(1))
        number = int(match.group(2))

        if current_chapter is not None and chapter != current_chapter:
            add_issue(
                issues,
                block,
                "公式编号",
                f"公式编号章节号与当前位置不一致：当前章节为第{current_chapter}章，编号为({chapter}-{number})。",
                f"建议改为“式({current_chapter}-{number})”或检查公式所在章节。"
            )

        expected = last_by_chapter.get(chapter, 0) + 1
        if number != expected:
            add_issue(
                issues,
                block,
                "公式编号",
                f"公式序号不连续：本章上一处为({chapter}-{last_by_chapter.get(chapter, 0)})，当前为({chapter}-{number})。",
                f"建议检查是否遗漏公式，或改为“式({chapter}-{expected})”。"
            )

        last_by_chapter[chapter] = number

    return issues


def check_format_blocks(blocks):
    issues = []
    issues.extend(check_heading_hierarchy(blocks))
    issues.extend(check_caption_sequence(blocks, "图", "图编号"))
    issues.extend(check_caption_sequence(blocks, "表", "表编号"))
    issues.extend(check_formula_numbering(blocks))
    return sorted(issues, key=lambda item: item["line"])


def render_issue(issue):
    st.markdown(f"**第{issue['line']}段 · {issue['type']}**")
    detail = f"{issue['msg']}（来源：{issue['source']}"
    if issue.get("style"):
        detail += f"，样式：{issue['style']}"
    detail += "）"
    st.warning(detail)
    st.code(issue["old"], language="text")
    st.caption(issue["suggestion"])


def format_page():
    render_page_shell("📐 格式校对", "支持上传 PDF/DOCX 生成 GB/T 7714 参考文献，并上传 Word 论文自动检查标题层级、图表编号、公式编号")

    tab1, tab2 = st.tabs(["📚 参考文献生成", "🔍 Word格式检查"])

    with tab1:
        st.subheader("上传文献自动生成参考文献")
        reference_files = st.file_uploader(
            "上传 PDF 或 Word 文档（支持多个文件）",
            type=["pdf", "docx", "doc"],
            accept_multiple_files=True,
            help="系统会抽取文档题名、作者、年份、期刊/学校等信息，并生成GB/T 7714-2015格式。"
        )

        if st.button("识别并生成参考文献", type="primary", use_container_width=True):
            if not reference_files:
                st.warning("请先上传PDF或Word文档。")
                return

            with st.spinner("正在识别文档信息并生成参考文献..."):
                result = generate_references_from_documents(reference_files)
                st.success("✅ 参考文献生成完成")
                st.code(result, language="text")
                st.download_button(
                    "📥 下载参考文献",
                    result,
                    file_name="参考文献_GBT7714.txt",
                    mime="text/plain",
                    use_container_width=True
                )

    with tab2:
        st.subheader("上传Word论文检查格式")
        paper_file = st.file_uploader(
            "上传论文Word文档（.docx）",
            type=["docx", "doc"],
            help="建议上传.docx文件；旧版.doc请先另存为.docx。"
        )

        if st.button("开始检查Word格式", type="primary", use_container_width=True):
            if not paper_file:
                st.warning("请先上传论文Word文档。")
                return
            if paper_file.name.lower().endswith(".doc"):
                st.warning("暂不支持旧版 .doc 文件，请先在 Word/WPS 中另存为 .docx 后上传。")
                return

            with st.spinner("正在读取Word文档并检查格式..."):
                try:
                    blocks = extract_docx_blocks(paper_file)
                except Exception as exc:
                    st.error(f"Word文档解析失败：{exc}")
                    return

                if not blocks:
                    st.warning("未从文档中提取到可检查的正文内容。")
                    return

                issues = check_format_blocks(blocks)
                st.session_state.format_issues = issues
                st.session_state.format_blocks = blocks

        if "format_issues" in st.session_state and "format_blocks" in st.session_state:
            issues = st.session_state.format_issues
            blocks = st.session_state.format_blocks

            st.markdown("---")
            st.subheader("📋 格式检查结果")
            st.caption(f"已读取 {len(blocks)} 个正文/表格段落，发现 {len(issues)} 个疑似格式问题。")

            if not issues:
                st.success("✅ 未发现明显的标题层级、图表编号或公式编号问题")
            else:
                issue_types = sorted(set(issue["type"] for issue in issues))
                selected_types = st.multiselect(
                    "按问题类型筛选",
                    issue_types,
                    default=issue_types,
                    key="format_issue_type_filter"
                )
                for issue in issues:
                    if issue["type"] in selected_types:
                        render_issue(issue)

                report = "\n\n".join(
                    f"第{issue['line']}段 [{issue['type']}]\n"
                    f"原文：{issue['old']}\n"
                    f"问题：{issue['msg']}\n"
                    f"建议：{issue['suggestion']}"
                    for issue in issues
                )
                st.download_button(
                    "📥 下载格式检查报告",
                    report,
                    file_name="Word格式检查报告.txt",
                    mime="text/plain",
                    use_container_width=True
                )
